import os
import io
import fitz  # PyMuPDF
import docx
import logging
from typing import List, Tuple, Dict, Any

# 로깅 설정
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """다양한 포맷의 문서에서 텍스트를 추출하는 클래스"""
    
    def __init__(self):
        pass
    
    def process_documents(self, documents: List[Tuple[str, bytes]]) -> List[Dict[str, Any]]:
        """문서 리스트를 처리하여 텍스트를 추출합니다.
        
        Args:
            documents: (파일명, 바이너리 데이터) 튜플의 리스트
            
        Returns:
            List[Dict]: 처리된 문서 정보 (파일명, 텍스트 내용, 메타데이터 등)
        """
        processed_docs = []
        
        for filename, content in documents:
            logger.info(f"'{filename}' 문서 처리 중...")
            
            # 파일 확장자 추출
            _, ext = os.path.splitext(filename.lower())
            
            # 파일 형식에 따라 적절한 처리기 선택
            extracted_text = ""
            metadata = {"filename": filename, "format": ext[1:] if ext else "unknown"}
            
            try:
                if ext == '.pdf':
                    extracted_text, meta = self._process_pdf(content)
                    metadata.update(meta)
                elif ext == '.docx':
                    extracted_text = self._process_docx(content)
                elif ext == '.txt' or ext == '.md':
                    extracted_text = content.decode('utf-8', errors='replace')
                else:
                    logger.warning(f"지원되지 않는 파일 형식: {ext}")
                    extracted_text = f"지원되지 않는 파일 형식: {ext}"
                
                processed_docs.append({
                    "filename": filename,
                    "text": extracted_text,
                    "metadata": metadata
                })
                
                logger.info(f"'{filename}' 처리 완료. 추출된 텍스트 길이: {len(extracted_text)} 자")
                
            except Exception as e:
                logger.error(f"'{filename}' 처리 중 오류 발생: {str(e)}")
                processed_docs.append({
                    "filename": filename,
                    "text": f"오류: {str(e)}",
                    "metadata": metadata,
                    "error": True
                })
        
        return processed_docs
    
    def _process_pdf(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """PDF 파일에서 텍스트를 추출합니다.
        
        Args:
            content: PDF 파일의 바이너리 데이터
            
        Returns:
            tuple: (추출된 텍스트, 메타데이터 딕셔너리)
        """
        # BytesIO를 사용하여 메모리 상의 파일로 변환
        pdf_stream = io.BytesIO(content)
        
        try:
            # PyMuPDF로 PDF 열기
            pdf_document = fitz.open(stream=pdf_stream, filetype="pdf")
            
            # 메타데이터 추출
            metadata = {
                "page_count": len(pdf_document),
                "title": pdf_document.metadata.get("title", ""),
                "author": pdf_document.metadata.get("author", ""),
                "subject": pdf_document.metadata.get("subject", ""),
                "keywords": pdf_document.metadata.get("keywords", "")
            }
            
            # 각 페이지에서 텍스트 추출
            all_text = []
            for page_num, page in enumerate(pdf_document):
                text = page.get_text()
                # 페이지 번호와 함께 텍스트 추가
                all_text.append(f"[페이지 {page_num + 1}]\n{text}\n")
            
            # PDF 닫기
            pdf_document.close()
            
            return "\n".join(all_text), metadata
            
        except Exception as e:
            logger.error(f"PDF 처리 중 오류: {str(e)}")
            return f"PDF 처리 오류: {str(e)}", {"error": str(e)}
    
    def _process_docx(self, content: bytes) -> str:
        """DOCX 파일에서 텍스트를 추출합니다.
        
        Args:
            content: DOCX 파일의 바이너리 데이터
            
        Returns:
            str: 추출된 텍스트
        """
        docx_stream = io.BytesIO(content)
        
        try:
            # python-docx로 파일 열기
            doc = docx.Document(docx_stream)
            
            # 문단 텍스트 추출
            paragraphs = [para.text for para in doc.paragraphs]
            
            # 표 내용도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    paragraphs.append(" | ".join(row_text))
            
            return "\n".join(paragraphs)
            
        except Exception as e:
            logger.error(f"DOCX 처리 중 오류: {str(e)}")
            return f"DOCX 처리 오류: {str(e)}"

if __name__ == "__main__":
    # 간단한 테스트 코드
    from document_fetcher import DocumentFetcher
    
    fetcher = DocumentFetcher()
    processor = DocumentProcessor()
    
    # 문서 가져오기
    documents = fetcher.fetch_documents()
    
    # 문서 처리
    processed_docs = processor.process_documents(documents)
    
    # 결과 확인
    for doc in processed_docs:
        print(f"파일: {doc['filename']}")
        print(f"텍스트 길이: {len(doc['text'])} 자")
        print(f"텍스트 미리보기: {doc['text'][:150]}...")
        print("-" * 50) 